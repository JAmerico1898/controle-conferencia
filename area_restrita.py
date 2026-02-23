"""
Área Restrita — Gestão da Conferência, Inscrições e Dashboards
"""

import streamlit as st
import pandas as pd
import io
from datetime import datetime, date, timedelta
from config import (
    MESES_CONFERENCIA, PREDIOS, COMBINACOES_OCUPACAO,
    ESTADOS_BR, DISCIPULADOS, DATAS_CHEGADA, COLUNAS_PLANILHA,
)
from sheets import (
    carregar_controle, obter_conferencia_ativa,
    abrir_conferencia, fechar_conferencia,
    listar_conferencias, carregar_inscricoes,
    carregar_inscricoes_ativas, contar_vagas,
    cancelar_inscricao, atualizar_inscricao,
    obter_ocupacao_conferencia, calcular_quarto_cama,
)
from auth import logout


def exibir_area_restrita():
    col_title, col_logout = st.columns([5, 1])
    with col_title:
        st.markdown("## ⚙️ Área Administrativa")
        st.caption(f"Logado como: **{st.session_state.get('usuario', '')}**")
    with col_logout:
        st.write(""); st.write("")
        if st.button("🚪 Sair"):
            logout()

    tab_gestao, tab_inscricoes, tab_dash = st.tabs([
        "📋 Gestão da Conferência", "👥 Inscrições", "📊 Dashboards",
    ])

    with tab_gestao:
        _gestao_conferencia()
    with tab_inscricoes:
        _gestao_inscricoes()
    with tab_dash:
        _dashboards()


# ──────────────────────────────────────────────
# GESTÃO DA CONFERÊNCIA
# ──────────────────────────────────────────────

def _gestao_conferencia():
    conf_ativa = obter_conferencia_ativa()

    if conf_ativa:
        dt_ini = _fmt_data(conf_ativa.get("data_inicio", ""))
        dt_fim = _fmt_data(conf_ativa.get("data_fim", ""))
        st.success(
            f"🟢 Conferência ativa: **{conf_ativa['mes']}/{conf_ativa['ano']}** "
            f"(inscrições de {dt_ini} a {dt_fim}) — "
            f"Ocupação: {conf_ativa.get('ocupacao', 'N/A')}"
        )
        if st.button("🔴 Fechar Inscrições Agora", type="primary"):
            fechar_conferencia()
            st.success("Inscrições encerradas.")
            st.rerun()
    else:
        st.info("Nenhuma conferência ativa no momento.")

    st.divider()

    # Abrir nova conferência (Issue #2: sem campo ano separado; Issue #3: com ocupação)
    st.subheader("Abrir Inscrições para Nova Conferência")

    col1, col2 = st.columns(2)
    with col1:
        mes_opcoes = {v: k for k, v in MESES_CONFERENCIA.items()}
        mes_selecionado = st.selectbox("Mês da Conferência", list(mes_opcoes.keys()), key="adm_mes")
        data_inicio = st.date_input("Data de início das inscrições",
            value=date.today(), format="DD/MM/YYYY", key="adm_dt_ini")
    with col2:
        data_fim = st.date_input("Data de fim das inscrições",
            value=date.today() + timedelta(days=10), format="DD/MM/YYYY", key="adm_dt_fim")
        ocupacao = st.selectbox(
            "Ocupação dos alojamentos",
            list(COMBINACOES_OCUPACAO.keys()),
            key="adm_ocupacao",
        )

    if st.button("✅ Abrir Inscrições", type="primary"):
        if conf_ativa:
            st.error("Já existe conferência ativa. Feche-a antes.")
            return
        if data_fim <= data_inicio:
            st.error("Data de fim deve ser posterior à data de início.")
            return

        mes_num = mes_opcoes[mes_selecionado]
        ano = data_inicio.year
        try:
            nome_aba = abrir_conferencia(
                mes_num, int(ano),
                data_inicio.strftime("%Y-%m-%d"),
                data_fim.strftime("%Y-%m-%d"),
                ocupacao,
            )
            st.success(f"Inscrições abertas para **{mes_selecionado}/{ano}**")
            st.rerun()
        except Exception as e:
            st.error(f"Erro: {e}")

    st.divider()

    # Histórico (Issue #1: formato dd-mm-aaaa)
    st.subheader("Histórico de Conferências")
    df_ctrl = carregar_controle()
    if df_ctrl.empty:
        st.caption("Nenhuma conferência registrada.")
    else:
        df_display = df_ctrl.copy()
        for col in ["data_inicio", "data_fim"]:
            if col in df_display.columns:
                df_display[col] = df_display[col].apply(_fmt_data)
        st.dataframe(df_display, use_container_width=True, hide_index=True)


# ──────────────────────────────────────────────
# GESTÃO DE INSCRIÇÕES
# ──────────────────────────────────────────────

def _gestao_inscricoes():
    nome_aba = _selecionar_conferencia()
    if not nome_aba:
        return

    df = carregar_inscricoes(nome_aba)
    df_ativos = df[df["status"].astype(str).str.upper() != "CANCELADO"] if not df.empty else df

    st.markdown(f"**Inscritos ativos:** {len(df_ativos)} | **Cancelados:** {len(df) - len(df_ativos)}")

    if df_ativos.empty:
        st.info("Nenhuma inscrição registrada.")
        return

    # Issue #4, #5, #6: Colunas alojamento, quarto, cama
    cols_display = ["codigo_inscricao", "nome", "genero", "cidade", "estado",
                    "discipulado", "alojamento", "quarto", "cama", "tipo_cama",
                    "data_chegada", "email"]
    cols_present = [c for c in cols_display if c in df_ativos.columns]
    st.dataframe(df_ativos[cols_present], use_container_width=True, hide_index=True)

    # Exportação
    st.divider()
    col_exp1, col_exp2 = st.columns(2)
    with col_exp1:
        csv = df_ativos.to_csv(index=False).encode("utf-8")
        st.download_button("📥 Exportar CSV", data=csv, file_name=f"inscritos_{nome_aba}.csv", mime="text/csv")
    with col_exp2:
        try:
            buffer = io.BytesIO()
            df_ativos.to_excel(buffer, index=False, engine="openpyxl")
            st.download_button(
                "📥 Exportar Excel", data=buffer.getvalue(),
                file_name=f"inscritos_{nome_aba}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        except ImportError:
            st.caption("Instale openpyxl para exportar Excel.")

    st.divider()

    # Issue #7: Relatórios por prédio
    st.subheader("📄 Relatórios por Alojamento")
    _gerar_relatorios_predios(df_ativos, nome_aba)

    st.divider()

    # Cancelar inscrição
    st.subheader("Cancelar Inscrição")
    nomes_ativos = df_ativos["nome"].tolist() if not df_ativos.empty else []
    if nomes_ativos:
        nome_cancelar = st.selectbox("Selecione o inscrito", [""] + nomes_ativos, key="adm_cancel_nome")
        if nome_cancelar and st.button("⚠️ Cancelar esta inscrição", key="adm_btn_cancel"):
            cancelar_inscricao(nome_aba, nome_cancelar)
            st.success(f"Inscrição de **{nome_cancelar}** cancelada.")
            st.rerun()

    st.divider()

    # Editar inscrição (inclui mover para Prédio Extra)
    st.subheader("Editar Inscrição")
    if nomes_ativos:
        nome_editar = st.selectbox("Selecione o inscrito para editar", [""] + nomes_ativos, key="adm_edit_nome")
        if nome_editar:
            row = df_ativos[df_ativos["nome"] == nome_editar].iloc[0].to_dict()
            _formulario_edicao(nome_aba, row)


def _formulario_edicao(nome_aba: str, row: dict):
    with st.expander("✏️ Editando inscrição", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            nome = st.text_input("Nome", value=row.get("nome", ""), key="edit_nome")
            genero = st.selectbox("Gênero", ["Masculino", "Feminino"],
                index=0 if row.get("genero") == "Masculino" else 1, key="edit_genero")
            cidade = st.text_input("Cidade", value=row.get("cidade", ""), key="edit_cidade")
            estado_idx = ESTADOS_BR.index(row.get("estado", "RJ")) if row.get("estado") in ESTADOS_BR else 0
            estado = st.selectbox("Estado", ESTADOS_BR, index=estado_idx, key="edit_estado")
            disc_idx = DISCIPULADOS.index(row.get("discipulado")) if row.get("discipulado") in DISCIPULADOS else 0
            discipulado = st.selectbox("Discipulado", DISCIPULADOS, index=disc_idx, key="edit_disc")
        with col2:
            # Issue #4: alojamento aceita Prédio Antigo, Prédio Novo, Prédio Extra, Não
            aloj_opcoes = ["Não", "Prédio Antigo", "Prédio Novo", "Prédio Extra"]
            aloj_atual = row.get("alojamento", "Não")
            aloj_idx = aloj_opcoes.index(aloj_atual) if aloj_atual in aloj_opcoes else 0
            alojamento = st.selectbox("Alojamento", aloj_opcoes, index=aloj_idx, key="edit_aloj")

            tipo_cama_opts = ["Baixo", "Cima", "N/A"]
            tc_atual = row.get("tipo_cama", "N/A")
            tc_idx = tipo_cama_opts.index(tc_atual) if tc_atual in tipo_cama_opts else 2
            tipo_cama = st.selectbox("Tipo de cama", tipo_cama_opts, index=tc_idx, key="edit_cama")

            quarto = st.text_input("Quarto", value=str(row.get("quarto", "")), key="edit_quarto")
            cama = st.text_input("Cama", value=str(row.get("cama", "")), key="edit_cama_num")
            email = st.text_input("Email", value=row.get("email", ""), key="edit_email")

        if st.button("💾 Salvar Alterações", type="primary", key="btn_salvar_edicao"):
            dados_atualizados = {
                "codigo_inscricao": row.get("codigo_inscricao"),
                "nome": nome, "genero": genero, "cidade": cidade, "estado": estado,
                "discipulado": discipulado,
                "alojamento": alojamento,
                "quarto": quarto if alojamento != "Não" else "N/A",
                "cama": cama if alojamento != "Não" else "N/A",
                "tipo_cama": tipo_cama if alojamento != "Não" else "N/A",
                "data_chegada": row.get("data_chegada", ""),
                "jantar_sexta": row.get("jantar_sexta", ""),
                "almoco_sabado": row.get("almoco_sabado", ""),
                "jantar_sabado": row.get("jantar_sabado", ""),
                "cafe_domingo": "Sim" if alojamento != "Não" else "Não",
                "lanche_domingo": row.get("lanche_domingo", ""),
                "email": email,
                "data_inscricao": row.get("data_inscricao", ""),
                "status": row.get("status", "Ativo"),
                "data_cancelamento": row.get("data_cancelamento", ""),
            }
            try:
                atualizar_inscricao(nome_aba, row["nome"], dados_atualizados)
                st.success("Inscrição atualizada!")
                st.rerun()
            except Exception as e:
                st.error(f"Erro: {e}")


# ──────────────────────────────────────────────
# RELATÓRIOS DOCX (Issue #7)
# ──────────────────────────────────────────────

def _gerar_relatorios_predios(df: pd.DataFrame, nome_aba: str):
    """Gera relatórios docx por prédio."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        st.warning("Instale python-docx para gerar relatórios. Adicione 'python-docx' ao requirements.txt.")
        return

    col1, col2, col3 = st.columns(3)

    # Prédio Antigo
    with col1:
        df_antigo = df[df["alojamento"].astype(str) == "Prédio Antigo"].copy()
        if not df_antigo.empty:
            df_antigo = df_antigo.sort_values("nome")
            doc = Document()
            doc.add_heading(f"Relatório — Prédio Antigo", level=1)
            doc.add_paragraph(f"Conferência: {nome_aba} | Total: {len(df_antigo)} aluno(a)s")
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Nome do Aluno(a)"
            hdr[1].text = "Tipo de Cama"
            hdr[2].text = "Nº Cama"
            for _, r in df_antigo.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(r.get("nome", ""))
                row_cells[1].text = f"Cama de {str(r.get('tipo_cama', '')).lower()}"
                row_cells[2].text = str(r.get("cama", ""))
            buf = io.BytesIO()
            doc.save(buf)
            st.download_button("📄 Prédio Antigo", data=buf.getvalue(),
                file_name=f"relatorio_predio_antigo_{nome_aba}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.caption("Nenhum inscrito no Prédio Antigo.")

    # Prédio Novo
    with col2:
        df_novo = df[df["alojamento"].astype(str) == "Prédio Novo"].copy()
        if not df_novo.empty:
            df_novo = df_novo.sort_values("nome")
            doc = Document()
            doc.add_heading(f"Relatório — Prédio Novo", level=1)
            doc.add_paragraph(f"Conferência: {nome_aba} | Total: {len(df_novo)} aluno(a)s")
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Nome do Aluno(a)"
            hdr[1].text = "Quarto"
            hdr[2].text = "Tipo de Cama"
            hdr[3].text = "Nº Cama"
            for _, r in df_novo.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(r.get("nome", ""))
                row_cells[1].text = str(r.get("quarto", ""))
                row_cells[2].text = f"Cama de {str(r.get('tipo_cama', '')).lower()}"
                row_cells[3].text = str(r.get("cama", ""))
            buf = io.BytesIO()
            doc.save(buf)
            st.download_button("📄 Prédio Novo", data=buf.getvalue(),
                file_name=f"relatorio_predio_novo_{nome_aba}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.caption("Nenhum inscrito no Prédio Novo.")

    # Prédio Extra
    with col3:
        df_extra = df[df["alojamento"].astype(str) == "Prédio Extra"].copy()
        if not df_extra.empty:
            df_extra = df_extra.sort_values("nome")
            doc = Document()
            doc.add_heading(f"Relatório — Prédio Extra", level=1)
            doc.add_paragraph(f"Conferência: {nome_aba} | Total: {len(df_extra)} aluno(a)s")
            table = doc.add_table(rows=1, cols=1)
            table.style = "Light Grid Accent 1"
            table.rows[0].cells[0].text = "Nome do Aluno(a)"
            for _, r in df_extra.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(r.get("nome", ""))
            buf = io.BytesIO()
            doc.save(buf)
            st.download_button("📄 Prédio Extra", data=buf.getvalue(),
                file_name=f"relatorio_predio_extra_{nome_aba}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.caption("Nenhum inscrito no Prédio Extra.")


# ──────────────────────────────────────────────
# DASHBOARDS
# ──────────────────────────────────────────────

def _dashboards():
    nome_aba = _selecionar_conferencia(key_suffix="dash")
    if not nome_aba:
        return

    df = carregar_inscricoes_ativas(nome_aba)
    if df.empty:
        st.info("Nenhuma inscrição registrada.")
        return

    # Resumo Geral
    st.subheader("📊 Resumo Geral")
    total = len(df)
    alojados = len(df[df["alojamento"].astype(str) != "Não"])
    nao_alojados = total - alojados

    col1, col2, col3 = st.columns(3)
    col1.metric("Total de Inscritos", total)
    col2.metric("Alojados", alojados)
    col3.metric("Não Alojados", nao_alojados)

    # Ocupação por prédio
    vagas = contar_vagas(nome_aba)
    ocupacao = obter_ocupacao_conferencia(nome_aba)
    st.markdown("#### Ocupação dos Alojamentos")

    for predio_nome in ["Prédio Antigo", "Prédio Novo"]:
        if predio_nome in vagas:
            vg = vagas[predio_nome]
            cap = PREDIOS[predio_nome]
            genero_label = ""
            for g, p in ocupacao.items():
                if p == predio_nome:
                    genero_label = g
            label = f"{predio_nome} ({genero_label})" if genero_label else predio_nome
            col_a, col_b = st.columns(2)
            with col_a:
                st.markdown(f"**{label} — Baixo**")
                pct = vg["baixo_ocupado"] / cap["baixo"] if cap["baixo"] > 0 else 0
                st.progress(min(pct, 1.0), text=f"{vg['baixo_ocupado']}/{cap['baixo']}")
            with col_b:
                st.markdown(f"**{label} — Cima**")
                pct = vg["cima_ocupado"] / cap["cima"] if cap["cima"] > 0 else 0
                st.progress(min(pct, 1.0), text=f"{vg['cima_ocupado']}/{cap['cima']}")

    st.divider()

    # Por Gênero
    st.subheader("👤 Inscritos por Gênero")
    st.bar_chart(df["genero"].value_counts())

    st.divider()

    # Por Discipulado
    st.subheader("📖 Inscritos por Discipulado")
    st.bar_chart(df["discipulado"].value_counts())

    st.divider()

    # Por Estado
    st.subheader("🗺️ Inscritos por Estado")
    st.bar_chart(df["estado"].value_counts())
    with st.expander("Detalhamento por Cidade"):
        cidade_counts = df.groupby(["estado", "cidade"]).size().reset_index(name="inscritos")
        st.dataframe(cidade_counts.sort_values(["estado", "inscritos"], ascending=[True, False]),
                     use_container_width=True, hide_index=True)

    st.divider()

    # Refeições
    st.subheader("🍽️ Contagem de Refeições")
    refeicoes_map = {
        "Jantar de Sexta": "jantar_sexta",
        "Almoço de Sábado": "almoco_sabado",
        "Jantar de Sábado": "jantar_sabado",
        "Café da Manhã de Domingo": "cafe_domingo",
        "Lanche de Domingo": "lanche_domingo",
    }
    ref_data = []
    for nome_ref, col in refeicoes_map.items():
        qtd = len(df[df[col].astype(str).str.upper() == "SIM"]) if col in df.columns else 0
        ref_data.append({"Refeição": nome_ref, "Quantidade": qtd})
    df_ref = pd.DataFrame(ref_data)
    st.dataframe(df_ref, use_container_width=True, hide_index=True)
    st.bar_chart(df_ref.set_index("Refeição"))


# ──────────────────────────────────────────────
# UTILIDADES
# ──────────────────────────────────────────────

def _selecionar_conferencia(key_suffix: str = "insc") -> str | None:
    conf_ativa = obter_conferencia_ativa()
    todas = listar_conferencias()
    if not todas and not conf_ativa:
        st.info("Nenhuma conferência registrada.")
        return None
    opcoes = []
    if conf_ativa:
        aba_ativa = conf_ativa["nome_aba"]
        opcoes.append(f"🟢 {aba_ativa} (ativa)")
        for t in todas:
            if t != aba_ativa:
                opcoes.append(t)
    else:
        opcoes = todas
    if not opcoes:
        return None
    selecionada = st.selectbox("Conferência", opcoes, index=0, key=f"sel_conf_{key_suffix}")
    return selecionada.replace("🟢 ", "").replace(" (ativa)", "").strip()


def _fmt_data(data_str: str) -> str:
    """Formata YYYY-MM-DD para DD/MM/YYYY (Issue #1)."""
    try:
        dt = datetime.strptime(str(data_str), "%Y-%m-%d")
        return dt.strftime("%d/%m/%Y")
    except Exception:
        return str(data_str)
